import re
from typing import Any, Dict, List

from docx import Document


Q_HEADER_RE = re.compile(r"^(Q\d{2,3})\b")


def _clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _normalize(text: str) -> str:
    text = text.strip().lower()
    text = re.sub(r"[\u2013\u2014\-]", "-", text)
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _paragraphs(docx_path: str) -> List[str]:
    """Extract all text from document in order, including table cells."""
    from docx.oxml import parse_xml
    from docx.table import _Cell
    
    doc = Document(docx_path)
    out = []
    
    # Process document in element order using body
    for elem in doc.element.body:
        # Paragraph
        if elem.tag.endswith('}p'):
            p = None
            for para in doc.paragraphs:
                if para._element is elem:
                    p = para
                    break
            if p:
                t = _clean(p.text)
                if t:
                    out.append(t)
        # Table
        elif elem.tag.endswith('}tbl'):
            for table in doc.tables:
                if table._tbl is elem:
                    for row in table.rows:
                        for cell in row.cells:
                            t = _clean(cell.text)
                            if t:
                                out.append(t)
                    break
    
    return out


def _find_marker(lines: List[str], marker: str, exact_start: bool = False) -> int:
    for i, line in enumerate(lines):
        if exact_start:
            if line.upper().startswith(marker.upper()):
                return i
        elif marker in line:
            return i
    return -1


def _split_q_blocks(lines: List[str]) -> Dict[str, List[str]]:
    blocks: Dict[str, List[str]] = {}
    current_q = None
    buffer: List[str] = []

    for line in lines:
        m = Q_HEADER_RE.match(line)
        if m:
            if current_q and buffer:
                blocks[current_q] = buffer
            current_q = m.group(1)
            buffer = [line]
            continue
        if current_q:
            buffer.append(line)

    if current_q and buffer:
        blocks[current_q] = buffer
    return blocks


def _parse_options(lines: List[str]) -> List[Dict[str, str]]:
    out = []
    for line in lines:
        m = re.match(r"^([A-Z])\.\s+(.*)$", line)
        if m:
            out.append({"key": m.group(1), "text": m.group(2).strip()})
    return out


def _parse_dropdown_groups(lines: List[str]) -> Dict[str, List[str]]:
    groups: Dict[str, List[str]] = {}
    current_label = ""
    for line in lines:
        m = re.match(r"^[\u25b8\-]\s*(.+):\s*$", line)
        if m:
            current_label = m.group(1).strip()
            groups[current_label] = []
            continue
        if current_label:
            if Q_HEADER_RE.match(line):
                break
            if re.match(r"^[\u25b8\-]\s+", line):
                current_label = ""
                continue
            if line.upper().startswith("EXPLANATION") or line.upper().startswith("ANSWER"):
                continue
            if line in {"OPTIONS:", "DROPDOWN OPTIONS:", "CORRECT SELECTIONS:", "CORRECT SEQUENCE:"}:
                continue
            groups[current_label].append(line)
    # remove empty/noisy entries
    cleaned: Dict[str, List[str]] = {}
    for k, vals in groups.items():
        vals = [v.strip() for v in vals if v.strip()]
        if vals:
            cleaned[k] = vals
    return cleaned


def _parse_available_values(lines: List[str]) -> List[str]:
    in_section = False
    values: List[str] = []

    for line in lines:
        # Match any "Available <word(s)>:" header line (values, segments, options, actions, etc.)
        if re.match(r"Available \w[\w\s]*:", line, re.IGNORECASE):
            right = line.split(":", 1)[1].strip()
            if right:
                # Inline slash-separated: "Available values: a / b / c"
                return [x.strip() for x in right.split("/") if x.strip()]
            # Multi-line bullet section follows
            in_section = True
            continue

        if in_section:
            stripped = line.strip()
            # bullet characters: •, -, \u2022
            if stripped and stripped[0] in ("\u2022", "\u2023", "-", "*", "\u25e6"):
                values.append(stripped.lstrip("\u2022\u2023-*\u25e6 ").strip())
            elif stripped == "":
                continue
            else:
                in_section = False

    return values


def _parse_statements(lines: List[str]) -> List[str]:
    out = []
    for line in lines:
        if re.match(r"^\d+\.\s+", line):
            out.append(line)
    return out


def _parse_yesno_table_statements(lines: List[str]) -> List[str]:
    """Extract YES/NO statements from table-style exports (non-numbered rows)."""
    start_idx = -1
    for i, line in enumerate(lines):
        if "for each statement" in line.lower():
            start_idx = i
            break
    if start_idx < 0:
        return []

    out: List[str] = []
    for line in lines[start_idx + 1 :]:
        t = line.strip()
        if not t:
            continue
        # Stop at obvious section boundaries.
        if t in {"OPTIONS:", "DROPDOWN OPTIONS:", "EXHIBIT:", "DETAILED ANSWERS"}:
            break
        if Q_HEADER_RE.match(t):
            break
        # Skip table header/noise tokens.
        if t.lower() in {"statement", "yes", "no"}:
            continue
        if t in {"○", "◯", "●"}:
            continue
        # Keep substantial statement lines.
        if len(t.split()) >= 5:
            out.append(t)
    return out


def _parse_select_count(lines: List[str]) -> int:
    for line in lines:
        m = re.search(r"select\s+(\d+)", line, flags=re.IGNORECASE)
        if m:
            return int(m.group(1))
    return 1


def _parse_question_block(block: List[str]) -> Dict[str, Any]:
    header = block[0]
    m = Q_HEADER_RE.match(header)
    qcode = m.group(1) if m else ""

    topic = ""
    tm = re.search(r"Topic\s+\d+", header)
    if tm:
        topic = tm.group(0)

    body = block[1:]
    options = _parse_options(body)
    dropdown_groups = _parse_dropdown_groups(body)
    available_values = _parse_available_values(body)
    statements = _parse_statements(body)
    if not statements:
        statements = _parse_yesno_table_statements(body)
    select_count = _parse_select_count(body)

    # Noise tokens that come from linearised YES/NO tables — never show in display text.
    _NOISE = {"statement", "yes", "no", "○", "◯", "●"}

    question_lines = []
    for line in body:
        if line in {"OPTIONS:", "DROPDOWN OPTIONS:", "EXHIBIT:", "DETAILED ANSWERS", "CORRECT SELECTIONS:", "CORRECT SEQUENCE:"}:
            continue
        if line.startswith("Available values:"):
            question_lines.append(line)
            continue
        # Drop bare noise tokens (table headers / radio circle chars).
        if line.strip().lower() in _NOISE:
            continue
        question_lines.append(line)

    return {
        "qcode": qcode,
        "topic": topic,
        "qtype": "UNKNOWN",
        "question_text": "\n".join(question_lines).strip(),
        "options": options,
        "dropdown_groups": dropdown_groups,
        "available_values": available_values,
        "statements": statements,
        "select_count": select_count,
        "correct_answer": {},
        "explanation": "",
    }


def _parse_detailed_answer_block(block: List[str]) -> Dict[str, Any]:
    out: Dict[str, Any] = {
        "qtype": "UNKNOWN",
        "correct_answer": {},
        "explanation": "",
    }

    type_line = next((l for l in block if l.startswith("Type:")), "")
    if type_line:
        out["qtype"] = type_line.split(":", 1)[1].strip().upper()

    # Handle ANSWER: lines (for MC, MULTI, YESNO)
    answer_line = next((l for l in block if l.startswith("ANSWER:")), "")
    if answer_line:
        value = answer_line.split(":", 1)[1].strip()
        out["correct_answer"] = {
            "mode": "answer",
            "value": value,
        }

    # Handle ANSWERS: lines (for YESNO with multiple statements)
    answers_idx = -1
    for i, line in enumerate(block):
        if line.startswith("ANSWERS:"):
            answers_idx = i
            break
    if answers_idx >= 0:
        items = []
        for line in block[answers_idx + 1 :]:
            if line.startswith("EXPLANATION:"):
                break
            if "→" in line and any(yn in line for yn in [" Yes", " No"]):
                parts = line.split("→")
                if len(parts) == 2:
                    stmt = parts[0].strip()
                    answer_val = parts[1].strip()
                    items.append({"label": stmt, "value": answer_val})
        if items:
            out["correct_answer"] = {
                "mode": "items",
                "items": items,
                "ordered": False,
            }

    # Handle CORRECT SELECTIONS: / CORRECT SEQUENCE:
    if "CORRECT SELECTIONS:" in block or "CORRECT SEQUENCE:" in block:
        marker = (
            "CORRECT SELECTIONS:"
            if "CORRECT SELECTIONS:" in block
            else "CORRECT SEQUENCE:"
        )
        start = block.index(marker) + 1
        items = []
        for line in block[start:]:
            if line.startswith("EXPLANATION:"):
                break
            if line.startswith("▸"):
                if ":" in line:
                    left, right = line[1:].split(":", 1)
                    items.append({"label": left.strip(), "value": right.strip()})
                else:
                    items.append({"label": line[1:].strip(), "value": ""})
            elif re.match(r"^(Blank|Step)\s+\d+", line) and ":" in line:
                left, right = line.split(":", 1)
                items.append({"label": left.strip(), "value": right.strip()})
        out["correct_answer"] = {
            "mode": "items",
            "items": items,
            "ordered": marker == "CORRECT SEQUENCE:",
        }

    exp_idx = -1
    for i, line in enumerate(block):
        if line.startswith("EXPLANATION:"):
            exp_idx = i
            break
    if exp_idx >= 0:
        out["explanation"] = "\n".join(block[exp_idx + 1 :]).strip()

    return out


def _parse_quick_answer_key(lines: List[str]) -> Dict[str, str]:
    key_map: Dict[str, str] = {}
    for line in lines:
        m = re.match(r"^(Q\d{2,3}):\s*(.+?)\s+[\u2013\u2014-]\s+Topic\s+\d+\s*$", line)
        if not m:
            continue
        qcode = m.group(1)
        value = m.group(2).strip()
        key_map[qcode] = value
    return key_map


def _infer_correct_answer_from_quick_key(question: Dict[str, Any], raw_value: str) -> Dict[str, Any]:
    raw = (raw_value or "").strip()
    if not raw:
        return {}

    # YES/NO statements
    statements = question.get("statements") or []
    if statements:
        yn_values = re.findall(r"\b(Yes|No)\b", raw, flags=re.IGNORECASE)
        if yn_values:
            items = []
            for i, statement in enumerate(statements):
                label = statement.replace(" YES / NO", "").strip()
                if i < len(yn_values):
                    items.append({"label": label, "value": yn_values[i].title()})
            if items:
                return {
                    "mode": "items",
                    "items": items,
                    "ordered": False,
                }

    # YES/NO without parsed statements yet: create generic statement labels.
    yn_values = re.findall(r"\b(Yes|No)\b", raw, flags=re.IGNORECASE)
    if yn_values:
        items = [
            {"label": f"Statement {i + 1}", "value": val.title()}
            for i, val in enumerate(yn_values)
        ]
        return {
            "mode": "items",
            "items": items,
            "ordered": False,
        }

    # Dropdown/hotspot answers separated by pipes map to dropdown labels in order.
    dropdown_groups = question.get("dropdown_groups") or {}
    if dropdown_groups:
        parts = [p.strip() for p in raw.split("|") if p.strip()]
        labels = list(dropdown_groups.keys())
        items = []
        for i, label in enumerate(labels):
            value = parts[i] if i < len(parts) else ""
            items.append({"label": label, "value": value})
        if items:
            return {
                "mode": "items",
                "items": items,
                "ordered": False,
            }

    # Option questions (single or multi)
    options = question.get("options") or []
    if options:
        text_to_key = {_normalize(o.get("text", "")): o.get("key", "") for o in options}
        norm_raw = _normalize(raw)
        if norm_raw in text_to_key and text_to_key[norm_raw]:
            return {
                "mode": "answer",
                "value": text_to_key[norm_raw],
            }

        if re.fullmatch(r"[A-Z]", raw):
            return {
                "mode": "answer",
                "value": raw,
            }
        if re.fullmatch(r"[A-Z](\s*,\s*[A-Z])+", raw):
            letters = [x.strip() for x in raw.split(",") if x.strip()]
            return {
                "mode": "answer",
                "value": "".join(letters),
            }

    # Generic pipe-separated items fallback.
    if "|" in raw:
        parts = [p.strip() for p in raw.split("|") if p.strip()]
        items = [{"label": f"Item {i + 1}", "value": p} for i, p in enumerate(parts)]
        if items:
            return {
                "mode": "items",
                "items": items,
                "ordered": False,
            }

    # Last fallback: keep as answer text.
    return {
        "mode": "answer",
        "value": raw,
    }


def parse_docx_questions(docx_path: str) -> List[Dict[str, Any]]:
    lines = _paragraphs(docx_path)

    part1_idx = _find_marker(lines, "PART 1", exact_start=True)
    part2_idx = _find_marker(lines, "PART 2", exact_start=True)
    detailed_idx = _find_marker(lines, "DETAILED ANSWERS", exact_start=True)

    if part1_idx < 0 or part2_idx < 0 or detailed_idx < 0:
        raise ValueError("Could not find PART 1/PART 2/DETAILED ANSWERS markers in DOCX.")

    q_lines = lines[part1_idx + 1 : part2_idx]
    quick_key_lines = lines[part2_idx + 1 : detailed_idx]
    detailed_lines = lines[detailed_idx + 1 :]

    question_blocks = _split_q_blocks(q_lines)
    detailed_blocks = _split_q_blocks(detailed_lines)
    quick_key = _parse_quick_answer_key(quick_key_lines)

    parsed_questions: List[Dict[str, Any]] = []
    all_qcodes = sorted(set(question_blocks.keys()) | set(detailed_blocks.keys()))
    for qcode in all_qcodes:
        source_block = question_blocks.get(qcode) or detailed_blocks.get(qcode) or []
        if not source_block:
            continue
        q = _parse_question_block(source_block)
        if qcode in detailed_blocks:
            d = _parse_detailed_answer_block(detailed_blocks[qcode])
            q["qtype"] = d.get("qtype", "UNKNOWN")
            q["correct_answer"] = d.get("correct_answer", {})
            q["explanation"] = d.get("explanation", "")

        if (not q.get("correct_answer")) and qcode in quick_key:
            q["correct_answer"] = _infer_correct_answer_from_quick_key(q, quick_key[qcode])

        # Reconcile YES/NO statement list against the quick answer key.
        # Questions whose statements live inside an exhibit image are hardcoded here in full.
        # When a qcode is present, its full statement list replaces whatever was parsed.
        _EXHIBIT_STATEMENTS: Dict[str, List[str]] = {
            "Q015": [
                "The design supports the technical requirements for redundancy (redundant if an Azure region fails).",
                "The design supports autoscaling.",
                "The design requires a manual configuration if an Azure region fails.",
            ],
        }
        if q.get("qtype") == "YES/NO" and qcode in quick_key:
            yn_vals = re.findall(r"\b(Yes|No)\b", quick_key[qcode], flags=re.IGNORECASE)
            stmts = _EXHIBIT_STATEMENTS[qcode] if qcode in _EXHIBIT_STATEMENTS else q.get("statements", [])
            if len(stmts) != len(yn_vals):
                # Trim to however many values we have answer keys for
                stmts = stmts[:len(yn_vals)]
            q["statements"] = stmts
            items = [
                {"label": s, "value": yn_vals[i].title()}
                for i, s in enumerate(stmts)
            ]
            q["correct_answer"] = {"mode": "items", "items": items, "ordered": False}

        parsed_questions.append(q)

    return parsed_questions


def evaluate_answer(question: Dict[str, Any], user_answer: Dict[str, Any]) -> Dict[str, Any]:
    correct = question.get("correct_answer", {}) or {}
    mode = correct.get("mode")

    if not correct or not mode:
        return {
            "is_correct": False,
            "feedback": "No answer key is available for this question yet.",
        }

    if not user_answer:
        return {
            "is_correct": False,
            "feedback": "No answer provided.",
        }

    # === MODE: 'answer' (MC, MULTI simple string answers) ===
    if mode == "answer":
        expected = (correct.get("value") or "").strip().upper()
        selected = (user_answer.get("selected_option") or "").strip().upper()
        selected_multi = [x.upper() for x in (user_answer.get("selected_options") or [])]
        text_answer = (user_answer.get("text_answer") or "").strip().upper()

        # Single letter or multi-letter sequence (A, AB, BD, etc.)
        if expected and expected[0].isalpha():
            if len(expected) == 1:
                # Single choice (MC)
                actual = selected or text_answer
                ok = actual == expected
                return {
                    "is_correct": ok,
                    "feedback": "" if ok else f"Expected: {expected}",
                }
            else:
                # Multi-choice (AB, BD, etc.) - order doesn't matter, just matching set
                expected_set = set(expected)
                actual_set = set(selected_multi) if selected_multi else set(text_answer)
                ok = actual_set == expected_set
                return {
                    "is_correct": ok,
                    "feedback": ""
                    if ok
                    else f"Expected: {expected} (any order). Got: {''.join(sorted(actual_set))}",
                }
        else:
            # Text answer (unlikely for MC)
            ok = _normalize(text_answer) == _normalize(expected)
            return {
                "is_correct": ok,
                "feedback": "" if ok else f"Expected: {expected}",
            }

    # === MODE: 'items' (HOTSPOT, DRAGDROP, YESNO) ===
    if mode == "items":
        expected_items = correct.get("items") or []
        user_items = user_answer.get("item_answers") or {}
        ordered = correct.get("ordered", False)

        checks = []
        all_ok = True
        for item in expected_items:
            label = item.get("label", "")
            expected_value = (item.get("value", "") or "").strip()
            actual_value = (user_items.get(label, "") or "").strip()
            
            # Normalize for comparison
            exp_norm = _normalize(expected_value)
            act_norm = _normalize(actual_value)
            
            ok = exp_norm == act_norm or actual_value.upper() == expected_value.upper()
            all_ok = all_ok and ok
            checks.append(
                {
                    "label": label,
                    "ok": ok,
                    "expected": expected_value,
                    "actual": actual_value,
                }
            )

        feedback_lines = []
        for check in checks:
            status = "✓" if check["ok"] else "✗"
            feedback_lines.append(
                f"{status} {check['label']}: {check['actual'] or '(empty)'}"
            )
            if not check["ok"]:
                feedback_lines.append(f"  Expected: {check['expected']}")

        return {
            "is_correct": all_ok,
            "feedback": "\n".join(feedback_lines) if feedback_lines else "No inputs provided.",
        }

    # === FALLBACK ===
    text_answer = (user_answer.get("text_answer") or "").strip()
    expected_text = str(correct)
    ok = _normalize(text_answer) == _normalize(expected_text)
    return {
        "is_correct": ok,
        "feedback": "" if ok else f"Expected: {expected_text}",
    }
